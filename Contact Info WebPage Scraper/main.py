

#Carter McCauley - 200576066
#Date: 04/06/2024
#COMP1112 - Python Automation
#Final Project - 10%


# --- Program name: Contact Gatherer ---
# This program is designed to receive an undefined amount of URLS local to barrie from the user.
#   Once received, the program will search each page given for contact information such as,
#   street address, postal code, phone number, and email address
#
#   The program will try to grab at least a phone number and email address if physical addresses
#   cannot be located.
#   Data gathered is then written to a review text file and results to a docx for ease of reading.


#import statements for libraries
#3rd party or native.
import requests
import urllib.request
from bs4 import BeautifulSoup
import re
import pyinputplus as pyip
import os
import time
import datetime
from docx import Document



#grab / return soup from website given
def getHtml(webSite):

    #try to open url
    try:
        #open web
        webUrl = urllib.request.urlopen(webSite)
        #get page info in bytes
        data = webUrl.read()
        #parse HTML
        soup = BeautifulSoup(data,'html.parser')

    #if open url throws error then return -1 to send error list instead.
    #we dont care about specific errors, only that we are unable to get contact data from the page.
    except:
        soup = -1

    #return valid soup or error code as -1.
    return soup


#get site name from website Url
def getSiteName(page):

    #hold site name
    siteName = ''

    #count slashes for domain name
    for i in range(8,len(page)):
        if(page[i] == '/'):
            #Slice the domain name from whole URL...
            siteName = page[8:i]

    #return site name 
    return siteName


#find postal code in split string, and return index of 
def postalCodeSearch(textSplit):

    #search for barrie postal codes / addresses
    postalRegex = re.compile(r'(L4N|L4M)')

    #if search fails then return -1 as error code.
    postalIndex = -1

    #grabs first postal code match..
    for i in range(len(textSplit)):
        #search for postal code in this lists element
        searchObject = postalRegex.search(textSplit[i])
        #if found a barrie 'L4N' or 'L4M' then save index.
        if(searchObject != None):
            postalIndex = i
            
    #return index of match
    return postalIndex


#search for phone number based off of postal code index. Should be near postal code so search is limited...
def phoneNumberSearch(textFormat,postalIndex):

    #phone number string
    phoneNumber = ''

    #first number regex without brackets
    numberRegex01 = re.compile(r'\d\d\d([.-]|\s)\d\d\d([.-]|\s)\d\d\d\d')
    #try bracket number for search if first doesn't work...
    numberRegex02 = re.compile(r'^[(]\d\d\d[)]$')

    #no postal code found on page, just grab phone number on page instead.
    if(postalIndex == -1):
        for i in range(len(textFormat)):

            #check for first email regex option
            if(numberRegex01.search(textFormat[i]) != None):
                phoneNumber = textFormat[i]
                break
            #check for second email regex option
            elif(numberRegex02.search(textFormat[i]) != None):
                phoneNumber = textFormat[i] + ' ' + textFormat[i+1]
                break
  
    #do normal search based off of postal code proximity
    else:
        #start checking for phone number on first element past end of postal code.
        for i in range(postalIndex+2,postalIndex+25):

            #check for first email regex option
            if(numberRegex01.search(textFormat[i]) != None):
                phoneNumber = textFormat[i]
                break
            #check for second email regex option
            elif(numberRegex02.search(textFormat[i]) != None):
                phoneNumber = textFormat[i] + ' ' + textFormat[i+1]
                break
            
    #return phone Number
    return phoneNumber


#find physical address in proximity to postal code index on page.
def addressSearch(textFormat,postalIndex):

    #string to append address text too
    addressString = ''

    #common street types to look for
    streetList = ['DRIVE', 'DR' 'ROAD', 'RD', 'STREET', 'ST', 'CRESCENT', 'CR','AVENUE', 'AVE','BOULEVARD','TRAIL','LANE','LN']

    #control variable to change state of for loop search once street type is found.
    control = True
    streetNumberIndex = -1

    #iterate down from postcal code end and find streetType
    for i in range(postalIndex,postalIndex-20,-1):
        #make element upper for easier reading
        element = textFormat[i].upper()

        #look for street type first, then transition to look for street number.
        if(control):
            #check street type list for street string
            for streetType in streetList:
                if(streetType in element):
                    #flip control and now look for street number
                    control = False

        #now look for next number and thats our street #
        if(element.isdecimal()):
            streetNumberIndex = i
    
    #create address string according to first number found after a street type
    for i in range(streetNumberIndex,postalIndex):
        addressString += textFormat[i] + ' '

    #return string of full address up to the postal code...
    return addressString


#get email based on postal Index, if not, then search whole page for an email address.
def emailSearch(textFormat,postalIndex):

    #variable for email
    email = ''
    #regex for email address
    regexObject = re.compile(r'[a-zA-Z0-9!#$%^&*().-_;:]{4,25}[@][a-z]{3,30}(.com|.com.|.ca|.ca.|.org|.net|.edu)')

    #no postal code index, so search whole page for an email...
    if(postalIndex == -1):
        #check after the postal code for 25 elements
        for i in range(len(textFormat)):

            if(regexObject.search(textFormat[i]) != None):
                email = textFormat[i]
                break

    #Postal code found so search according to proximity to postal code.
    else:
        #check after the postal code for 25 elements
        for i in range(postalIndex+2,len(textFormat)):

            if(regexObject.search(textFormat[i]) != None):
                email = textFormat[i]
                break

    #return email address if found...
    return email


#gather each portion of contact data from given site text
    #will search for postal code first, and use that to find street address/phone#/email
    #this is done b/c there may be many numbers / other address, but
    #we are interested in only barrie addresses
def search(page,text,textFormat):

    #variable to append page data to
    pageData = []    

    #page is good, so append GOOD tag.
    pageData.append('GOOD')
    
    #get site Name
    siteName = getSiteName(page)
    #domain always gets appended so append up here
    pageData.append(siteName)

    #find address string by finding postal code
        #return index of first part of postal code found
    postalIndex = postalCodeSearch(textFormat)

    #If no postal code found, then just search for a phone number on the whole page.
    if(postalIndex != -1):
        postCodeStart = textFormat[postalIndex]
        postCodeEnd = textFormat[postalIndex+1]
        #get postal code ready for list append.
        postalCode = postCodeStart + ' ' + postCodeEnd

        #get address string based off of postal code index.
        addressString = addressSearch(textFormat,postalIndex)

        #get phone number based off of postal code index.
        phoneNumber = phoneNumberSearch(textFormat,postalIndex)

        #get email address.
        email = emailSearch(textFormat, postalIndex)
    
        #append in order of:
            #street address
            #postal code
            #phone number
            #email address
        pageData.append(addressString)
        pageData.append(postalCode)
        pageData.append(phoneNumber)
        pageData.append(email)

    #postalcode returns -1, so search for a phonenumber and email since they are more unique
        #DO NOT look for street address because numbers / street types are too common, and junk data may be gathered by accident.
    else:
        #secondary search, postalIndex = -1, so just find a phone number on whole page.
        phoneNumber = phoneNumberSearch(textFormat,postalIndex)

        #get email address.
        email = emailSearch(textFormat, postalIndex)
        #append blanks for street address + postal code
        pageData.append('')
        pageData.append('')
        #add phone number to pageData
        pageData.append(phoneNumber)
        pageData.append(email)

    #return data acquired from page
    return pageData

    
    
#takes list of urls
#gathers data from each url in turn 
def searchMain(listOfUrls):

    #create list to store each pages list of contact info
    contactData = []

    #search each page
    for page in listOfUrls:

        #grab soup from page
        soup = getHtml(page)

        #if error is thrown when trying to open page, send error list with domain name.
        if(soup == -1):
            siteName = getSiteName(page)
            badPage = []
            badPage.append('ERROR')
            badPage.append(siteName)
            #send empty tags, so list remains standard length 
            badPage.append('')
            badPage.append('')
            badPage.append('')
            badPage.append('')
            contactData.append(badPage)

        #page is good, and can read data.
        else:
            #get text from html content
            text = soup.get_text()
            #split text by space ' '
            textFormat = text.split()

            #return list of pages contact data in order..
            pageList = search(page,text,textFormat)

            #append list of page data, into main contact Data list.
            contactData.append(pageList)
        
    #return list of lists which hold ordered contact information of page/biz
    return contactData


#grab as many links as the user wants to put into the program
def getLinksFromUser():

    #list of links to append to and return
    linkList = []

    #prompt user about how to input links
    print('')
    print('Please enter your URL as is. When you are finished, please enter "done"')
    #sleep so user sees this and understands steps to program
    time.sleep(3)

    #while loop so user can enter 'done' when finished giving links.
    control = False
    while(not control):
        #get user input
        #only accept either 'done' or word beginning with 'https://'
        userInput = pyip.inputStr(blockRegexes=[r'.*'],
                    allowRegexes=[r'(^done$|^https://)'],
                    prompt='Enter a URL: '
                                  )
        #user is done inputting links so end loop
        if(userInput == 'done'):
            control = True
        #append url to list
        else:
            linkList.append(userInput)

    #return the list of urls given by user
    return linkList


#see how many errors there were with the URL links given.
    #return list of domain names to give to report...
def getErrorCount(dataList):

    #check for amount of errors
    errorCountList = []
    
    #iterate lists to find amount of errors
    for page in dataList:
        if(page[0] == 'ERROR'):
            errorCountList.append(page[1])
    #return amount of errors
    return errorCountList

#check for no postal codes in data.
    #return list of domain names
def noPostalCodeCheck(dataList):
    #append domain name to list if no postal code
    noCodeList = []

    #iterate through data list
    for page in dataList:
        if(page[3] == ''):
            if(page[0] != 'ERROR'):
                noCodeList.append(page[1])
    #return count of NO postal codes
    return noCodeList


#write a summary of the information gathered into a text file for review
def writeSummary(dataList,programExecution):

    #returns path of current directoy of this file.
    pather = os.getcwd()
    new = pather.replace('\\','\\\\')

    #gather some information before wrting
        #how many links were received to look at.
        #how many links had errors when trying to gain access.
        #how many links are good.
        #how many links had no postal code to find.

    #length of master links
    lengthOfDataList = len(dataList)

    #call errorCount to get amount of errors.
    errorCountList = getErrorCount(dataList)

    #get how many links were good by subtracting errorCount from length of links
    goodLinks = lengthOfDataList - len(errorCountList)

    #call checkPostalCount to see how many links had no postal codes
    postalCodeList = noPostalCodeCheck(dataList)
    
    #open a text file to write too
    file = open(new+'\\\\TEST01.txt','w')

    #display header info in text file
    file.write('\n---Contact Gatherer Report Summary---\n')
    file.write('\n-')
    file.write(str(datetime.datetime.now()))
    file.write('')
    file.write('\n\nThis file is a summary of the results for your most recent use of the program Contact Gatherer.\n\n')
    file.write('\n----------\n')
    
    #display general information about errors, total links checked, and other info.
    file.write('Review:\n')
    file.write('----------\n')
    #show program execution time
    file.write('Program executed in: ')
    file.write(str(programExecution))
    #show how many sites were given to the program
    file.write('\n\nThe amount of websites given to the program: ')
    file.write(str(lengthOfDataList))
    file.write(' sites\n\n')
    #good links count
    file.write('Good Links: ')
    file.write(str(goodLinks))
    file.write('\n\n----------\n')
    #error section 
    file.write('Error Section:\n')
    file.write('----------\n')
    #show the amount of links that had an error
    file.write('Error Links: ')
    file.write(str(len(errorCountList)))
    file.write('\n')
    file.write('\n')
    #if errors found, then show the domain name that was bad
    if(len(errorCountList) > 0):
        for i in range(len(errorCountList)): 
            file.write('\t')
            file.write(errorCountList[i])
            file.write('\n\n')
    
    #display number of no postal codes
    file.write('Links that had no postal codes: ')
    file.write(str(len(postalCodeList)))
    file.write('\n\n')
    #display the site that had no postal code / physical address.
    if(len(postalCodeList) > 0):
        for i in range(len(postalCodeList)):
            file.write('\t')
            file.write(postalCodeList[i])
            file.write('\n')

    #display some general end page content
    file.write('\n\n----------\n')
    file.write('This report is brought to you by MC Technologies.\n')
    file.write('Copyright @ 2017-federal-patents.')
    
    #close file that was written too.
    file.close()
    

#this function will write the found contact data to a DOCX file.
def writeToDoc(dataList):

    #create new document in current directory.
    doc = Document()

    #Add a heading to the document
    doc.add_heading('Contact Gatherer Program Results',0)

    #add a timestamp for user
    doc.add_paragraph(str(datetime.datetime.now()))

    #add a descriptive paragraph
    doc.add_paragraph('''This report contains the findings of contact information for the given website links from the Contact Gatherer Program. Below will be listed the domain names that were given, and the results found according to each of them. Note: It is possible for some links to have failed and access have been denied. In this case, an Error will be shown for its "Status", and blank results for its other information.''')
    
    #write out headers for the data
    #6 headers / columns to write
    #Status - Domain name - Street Address - Postal Code - Phone # - Email Address
    headers = ['Status','Domain Name','Street Address','Postal Code','Phone Number','Email Address']

    #add header showcasing the results section within the docx.
    doc.add_heading('Results',1)

    #iterate through pages data and write them to docx file.
    for i in range(len(dataList)):

        #add blank paragraph for spacing
        doc.add_paragraph('')
        #add heading of page content.
        doc.add_heading(dataList[i][1],2)
        
        #display results for page data
        #display status for page 
        doc.add_paragraph('Status: ')
        doc.add_paragraph(dataList[i][0])
        #display street address for page
        doc.add_paragraph('Street Address:')
        doc.add_paragraph(dataList[i][2])
        #display postal code for page
        doc.add_paragraph('Postal Code:')
        doc.add_paragraph(dataList[i][3])
        #display phone number for page
        doc.add_paragraph('Phone Number:')
        doc.add_paragraph(dataList[i][4])
        #display email address for page
        doc.add_paragraph('Email Address:')
        doc.add_paragraph(dataList[i][5])

    #add some blank space
    doc.add_paragraph('')
    doc.add_paragraph('')
    #add information paragraph at bottom
    doc.add_paragraph('''Note for results containing no postal code: that a postal code was not found on the given web page therefore, it was too unreliable to grab a street address along with it. If this is the case then a phone number and email address will try to be grabbed.''')
    
    #save the doc after changes have been made..
    doc.save('TESTER1000.docx')


#simply welcome user, and define what they need to do for the program to work.
def programStart():

    #welcome user to the program
    print('Welcome to Contact Gatherer!\n')
    #sleep for readability
    time.sleep(2)
    print('This program is designed to find contact information about Barrie business\'.\n')
    #sleep for consistent output
    time.sleep(1)
    print('You will need to provide URL\'s for the program to use.\n')
    #sleep for conssitent output
    time.sleep(1)
    print('Please follow the instructions below.')
    #make program sleep to user can read text
    time.sleep(3)

#control main program flow
#able to take lists of local website Urls to grab contact info from all of them
def main():
    
    #welcome user and describe program
    programStart()
    
    #grab links from user
    pages = getLinksFromUser()

    #clock in start time of program.
    timeStart = time.time()

    #grab contact info from all the sites
        #Site name / business
        #Address
        #City
        #Phone number
        #Postal Code
        #email address
    #RETURN - list of lists. Each list within is dedicated for that page given.
    dataList = searchMain(pages)

    #get time of program end
    timeEnd = time.time()
    executionTime = timeEnd-timeStart

    #write summary to text file
    writeSummary(dataList,executionTime)

    #write data to excel file.....
    writeToDoc(dataList)

    #tell user program is finished for ease of mind. have a little fun with user.
    print('\nExecution finished...\n')
    print('That took about ',str(executionTime),' seconds...\n')
    print('Now thats impressive.\n')
    #sleep for a second for user experience
    time.sleep(2)
    print('Please check your directory for program results and review.')

    #END OF PROGRAM


#
#---PROGRAM START---
#
#run the program through main program control.
main()





